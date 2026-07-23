from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path("fixtures/Customer_Complaint_Metformin_MTF240821.docx")
BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
GRAY = "5B6573"

def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)

def cell_text(cell, text, bold=False, color=None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def details_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(2.05); cells[1].width = Inches(4.45)
        shade(cells[0], LIGHT_BLUE)
        cell_text(cells[0], label, bold=True, color=BLUE)
        cell_text(cells[1], value)
        for c in cells: set_cell_margins(c)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55); section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"; normal.font.size = Pt(9.5)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.04
for style_name, size in (("Heading 1", 15), ("Heading 2", 12)):
    st = doc.styles[style_name]; st.font.name = "Aptos Display"; st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(BLUE)
    st.paragraph_format.space_before = Pt(8); st.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CUSTOMER PRODUCT QUALITY COMPLAINT")
r.bold = True; r.font.name = "Aptos Display"; r.font.size = Pt(17); r.font.color.rgb = RGBColor.from_string(BLUE)
p.paragraph_format.space_after = Pt(3)
p = doc.add_paragraph("Demonstration document for QualiSphere AI Complaint Intake")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

details_table(doc, [
    ("Complaint reference", "MD-CC-2026-0718"),
    ("Date reported", "18 July 2026"),
    ("Complaint source", "Distributor"),
    ("Customer / reporter", "Priya Shah, Quality Lead — Medico Distribution"),
    ("Priority requested", "Urgent — product quality assessment requested"),
])

doc.add_heading("1. Product and batch identification", level=1)
details_table(doc, [
    ("Product name", "Metformin Hydrochloride Tablets 500 mg"),
    ("Dosage form", "Film-coated tablets in PVC/PVdC blister packs"),
    ("Batch / lot number", "MTF240821"),
    ("Manufacturing date", "21 August 2024"),
    ("Expiry date", "31 August 2027"),
    ("Quantity affected", "18 tablets from 2 blister strips"),
])

doc.add_heading("2. Complaint description", level=1)
doc.add_paragraph("Dear Quality Team,")
doc.add_paragraph(
    "We received Metformin Hydrochloride Tablets 500 mg, batch MTF240821, at our distribution warehouse. During incoming visual inspection, several tablets were found broken inside otherwise intact blister pockets. The outer cartons and shipper cases showed no visible damage.")
doc.add_paragraph(
    "A total of 18 tablets were affected across two packs. The broken tablets had visible fragments within the blister cavities. No product was dispensed to patients and no adverse event has been reported. We have segregated the affected packs and retained the remaining stock pending your direction.")
doc.add_paragraph(
    "Please investigate the potential cause of tablet breakage and advise whether this could affect other packs or batches. Photographs and retained samples are available on request.")
doc.add_paragraph("Regards,\nPriya Shah\nQuality Lead, Medico Distribution")

doc.add_heading("3. Customer observations", level=1)
for item in [
    "Breakage was observed inside intact blister pockets; no external transport damage was apparent.",
    "Complaint was identified before patient dispensing; no patient impact is known.",
    "Affected stock has been quarantined at the distributor warehouse.",
    "Customer requests investigation feedback and disposition guidance before release of remaining inventory.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4. QA intake notes", level=1)
p = doc.add_paragraph()
r = p.add_run("For QMS test use only. "); r.bold = True; r.font.color.rgb = RGBColor.from_string("9B1C1C")
p.add_run("This document contains fictional names, product details, and complaint information. It is not a real customer complaint and must not be used for GMP records or regulatory decisions.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("QualiSphere demo fixture | Fictional data | Customer complaint test document")
fr.font.name = "Aptos"; fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string(GRAY)

OUT.parent.mkdir(exist_ok=True)
doc.core_properties.title = "Customer Complaint — Metformin Hydrochloride Tablets"
doc.core_properties.subject = "QualiSphere test fixture"
doc.core_properties.author = "QualiSphere Demo"
doc.save(OUT)
print(OUT.resolve())
